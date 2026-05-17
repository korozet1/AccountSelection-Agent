from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


INPUT = Path("inputs/word_test_source.docx")
OUTPUT = Path("outputs/word_test/YXC-PM-03-Word 操作测试题-复核修正版.docx")


def get_or_add(parent, tag):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_run_fonts(run, size_pt=None, hidden=False, red=False, highlight_red=False):
    font = run.font
    font.name = "Times New Roman"
    if size_pt is not None:
        font.size = Pt(size_pt)
    if hidden:
        font.hidden = True
    if red:
        font.color.rgb = __import__("docx.shared").shared.RGBColor(255, 0, 0)
    if highlight_red:
        font.highlight_color = WD_COLOR_INDEX.RED

    r_pr = run._element.get_or_add_rPr()
    r_fonts = get_or_add(r_pr, "w:rFonts")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")


def set_paragraph_spacing(paragraph, size_pt=12):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    for run in paragraph.runs:
        set_run_fonts(run, size_pt=size_pt)


def set_cell_text_font(cell, size_pt=11):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_before = Pt(0)
        for run in paragraph.runs:
            set_run_fonts(run, size_pt=size_pt)


def set_table_percent_width(table):
    tbl_pr = table._tbl.tblPr
    tbl_w = get_or_add(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is not None:
        tbl_pr.remove(tbl_layout)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.allow_autofit = True

    col_count = len(table.columns)
    if col_count:
        pct_each = str(int(5000 / col_count))
        tbl_grid = table._tbl.tblGrid
        if tbl_grid is None:
            tbl_grid = OxmlElement("w:tblGrid")
            table._tbl.insert(0, tbl_grid)
        for child in list(tbl_grid):
            tbl_grid.remove(child)
        for _ in range(col_count):
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(int(9026 / col_count)))
            tbl_grid.append(grid_col)

        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = get_or_add(tc_pr, "w:tcW")
                tc_w.set(qn("w:w"), pct_each)
                tc_w.set(qn("w:type"), "pct")


def replace_paragraph_text_with_single_run(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
        return paragraph.runs[0]
    return paragraph.add_run(text)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def set_hidden_text_display_off(doc):
    settings = doc.settings.element
    for tag in ("w:displayHiddenText", "w:printHiddenText"):
        node = settings.find(qn(tag))
        if node is not None:
            settings.remove(node)
    display_hidden = OxmlElement("w:displayHiddenText")
    display_hidden.set(qn("w:val"), "false")
    settings.append(display_hidden)
    print_hidden = OxmlElement("w:printHiddenText")
    print_hidden.set(qn("w:val"), "false")
    settings.append(print_hidden)


def set_default_style_fonts(doc):
    for style in doc.styles:
        if getattr(style, "type", None) is None:
            continue
        try:
            style.font.name = "Times New Roman"
            if style.name == "Normal":
                style.font.size = Pt(12)
            r_pr = style.element.get_or_add_rPr()
            r_fonts = get_or_add(r_pr, "w:rFonts")
            r_fonts.set(qn("w:ascii"), "Times New Roman")
            r_fonts.set(qn("w:hAnsi"), "Times New Roman")
            r_fonts.set(qn("w:cs"), "Times New Roman")
            r_fonts.set(qn("w:eastAsia"), "宋体")
        except Exception:
            pass


def main():
    doc = Document(INPUT)
    set_default_style_fonts(doc)
    set_hidden_text_display_off(doc)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        header_p = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        header_run = replace_paragraph_text_with_single_run(header_p, "Final Report")
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_p.paragraph_format.space_before = Pt(0)
        header_p.paragraph_format.line_spacing = 1.5
        set_run_fonts(header_run, size_pt=12)

        footer_p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        footer_run = replace_paragraph_text_with_single_run(footer_p, "Page 1 of 78")
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.paragraph_format.space_before = Pt(0)
        footer_p.paragraph_format.line_spacing = 1.5
        set_run_fonts(footer_run, size_pt=12)

    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() == "Page 1 of 78":
            delete_paragraph(paragraph)
            continue
        set_paragraph_spacing(paragraph, size_pt=12)
        if "Total 78" in paragraph.text:
            for run in paragraph.runs:
                if "Total 78" in run.text:
                    set_run_fonts(run, size_pt=12, highlight_red=True)
                elif paragraph.text == "Total 78":
                    set_run_fonts(run, size_pt=12, highlight_red=True)
        if paragraph.text == "Total 78" and not paragraph.runs:
            run = paragraph.add_run("Total 78")
            set_run_fonts(run, size_pt=12, highlight_red=True)

        if "Shin Nippon Biomedical Laboratories, Ltd." in paragraph.text:
            for run in paragraph.runs:
                set_run_fonts(run, size_pt=12, hidden=True)

    for table in doc.tables:
        set_table_percent_width(table)
        for row in table.rows:
            row.height = None
            for cell in row.cells:
                set_cell_text_font(cell, size_pt=11)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
