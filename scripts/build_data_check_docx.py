from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT = Path("outputs/data_check/YXC-PM-04-数据核查测试题-完成版.docx")
START = Path("outputs/data_check/start_time.txt").read_text(encoding="utf-8-sig").strip()
END = datetime.now().astimezone().strftime("%Y-%m-%d %H:%M:%S %z")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def add_run(paragraph, text, bold=False, size=8.5, highlight=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")
    rfonts.set(qn("w:eastAsia"), "宋体")
    return run


def para(doc, text="", bold=False, size=8.5, align=None, before=0, after=1.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    if align is not None:
        p.alignment = align
    if text:
        add_run(p, text, bold=bold, size=size)
    return p


def paragraph_with_segments(doc, segments, size=8.5, after=1.5):
    p = para(doc, after=after)
    for text, highlight in segments:
        add_run(p, text, size=size, highlight=highlight)
    return p


def table(doc, rows, widths, highlight_cells=None, font_size=7.2, header_rows=1):
    highlight_cells = highlight_cells or set()
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = t.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_width(cell, widths[c_idx])
            if r_idx < header_rows:
                set_cell_shading(cell, "BFBFBF")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_run(
                p,
                text,
                bold=r_idx < header_rows,
                size=font_size,
                highlight=(r_idx, c_idx) in highlight_cells,
            )
    return t


def add_spacer(doc, height_pt=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(height_pt)
    p.paragraph_format.line_spacing = 1.0
    return p


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.top_margin = Cm(0.8)
    sec.bottom_margin = Cm(0.8)
    sec.left_margin = Cm(1.0)
    sec.right_margin = Cm(1.0)

    para(doc, f"测试开始时间：{START}    测试结束时间：{END}", size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)

    hdr = doc.add_table(rows=3, cols=2)
    hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr.autofit = False
    set_cell_width(hdr.cell(0, 0), 13.5)
    set_cell_width(hdr.cell(0, 1), 4.5)
    hdr.cell(0, 0).text = "STA/SMA/0824"
    hdr.cell(0, 1).text = "Révision : 3"
    hdr.cell(1, 0).merge(hdr.cell(1, 1))
    hdr.cell(1, 0).text = "REAL TIME STABILITY STUDY REPORT"
    hdr.cell(2, 0).merge(hdr.cell(2, 1))
    hdr.cell(2, 0).text = "TITLE: VIDAS 25 OH VITAMIN D / REF 30463\nASSOCIATED STABILITY PROTOCOL : PCS/SMA/0824"
    for row in hdr.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if "REAL TIME" in p.text or "Révision" in p.text else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.bold = True
                    r.font.name = "Arial"
                    r.font.size = Pt(9)
    add_spacer(doc, 2)

    table(
        doc,
        [
            ["", "PTB1", "PTB2", "PTB3"],
            ["Sample", "Expected values RFV", "Expected values RFV", "Expected values RFV"],
            ["S1", "1739 - 4520", "1265 - 3288", "1392 - 3618"],
        ],
        [2.2, 5.1, 5.1, 5.1],
        highlight_cells={(2, 1), (2, 2), (2, 3)},
        header_rows=1,
    )

    para(doc, "❖  Kit Control C1", bold=True, size=9.5, after=1)
    para(
        doc,
        "Targets were determined by values obtained on day 0 of this study and expected values ranges as target +/- 3 SD, "
        "SD determined using the between-lot precision profile based on 3 lots (PTB1, PTB2 & PTB3).",
        size=8.5,
    )
    table(
        doc,
        [
            ["", "PTB1", "PTB1", "PTB2", "PTB2", "PTB3", "PTB3"],
            ["Sample", "Target\n(ng/mL)", "Expected values\n(ng/mL)", "Target\n(ng/mL)", "Expected values\n(ng/mL)", "Target\n(ng/mL)", "Expected values\n(ng/mL)"],
            ["C1", "37.1", "33.0 - 41.2", "52.6", "48.3 - 56.9", "52.2", "47.9 - 56.5"],
        ],
        [1.5, 2.1, 3.2, 2.1, 3.2, 2.1, 3.2],
        highlight_cells={(2, 2)},
        header_rows=1,
    )
    paragraph_with_segments(
        doc,
        [
            ("Expected values were also calculated for each serum by the R&D laboratory with the VIDAS 25 OH Vitamin D test by using ", False),
            ("CV%", True),
            (" PRD specification for a second level of interpretation.", False),
        ],
    )
    table(
        doc,
        [
            ["", "PTB1", "PTB1", "PTB2", "PTB2", "PTB3", "PTB3"],
            ["Sample", "Target\n(ng/mL)", "Expected values\n(ng/mL)", "Target\n(ng/mL)", "Expected values\n(ng/mL)", "Target\n(ng/mL)", "Expected values\n(ng/mL)"],
            ["C1", "37.1", "25.9 - 48.2", "52.6", "32.1 - 73.1", "52.2", "31.8 - 72.6"],
        ],
        [1.5, 2.1, 3.2, 2.1, 3.2, 2.1, 3.2],
        highlight_cells={(1, 0), (1, 4), (2, 2), (2, 4), (2, 6)},
        header_rows=1,
    )

    para(doc, "❖  Samples", bold=True, size=9.5, after=1)
    para(
        doc,
        "Targets were determined by values obtained on day 0 of this study and expected values ranges as target +/- 3 SD, "
        "SD determined using the between-lot precision profile based on 3 lots (PTB1, PTB2 & PTB3).",
        size=8.5,
    )
    table(
        doc,
        [
            ["", "PTB1", "PTB1", "PTB2", "PTB2", "PTB3", "PTB3"],
            ["Sample", "Target\n(ng/mL)", "Expected values\n(ng/mL)", "Target\n(ng/mL)", "Expected values\n(ng/mL)", "Target\n(ng/mL)", "Expected values\n(ng/mL)"],
            ["SCI 1", "21.3", "17.3 - 25.4", "21.3", "17.2 - 25.4", "21.6", "17.5 - 25.7"],
            ["SCI 2", "33.3", "29.2 - 37.4", "34.1", "30.0 - 38.2", "34.2", "30.1 - 38.3"],
            ["SCI 3", "42.2", "38.0 - 46.4", "43.4", "39.2 - 47.6", "43.3", "39.1 - 47.5"],
            ["SCI 4", "64.9", "60.3 - 69.6", "66.1", "61.4 - 70.8", "65.3", "60.6 - 70.0"],
            ["SCI 5", "82.5", "77.2 - 87.8", "85.1", "79.7 - 90.5", "85.1", "79.7 - 90.5"],
        ],
        [1.5, 2.1, 3.2, 2.1, 3.2, 2.1, 3.2],
        highlight_cells={(2, 6), (3, 1), (3, 6), (4, 2), (4, 4), (4, 6), (5, 3), (5, 4), (6, 2), (6, 4), (6, 6)},
        header_rows=1,
    )
    paragraph_with_segments(
        doc,
        [
            ("Expected values were also calculated for each serum by the R&D laboratory with the VIDAS 25 OH Vitamin ", False),
            ("D", True),
            (" test by using CV% PRD specification ", False),
            ("for", True),
            (" a second level of interpretation.", False),
        ],
    )
    table(
        doc,
        [
            ["", "PTB1", "PTB1", "PTB2", "PTB2", "PTB3", "PTB3"],
            ["Sample", "Target\n(ng/mL)", "PRD Expected\nvalues (ng/mL)", "Target\n(ng/mL)", "PRD Expected\nvalues (ng/mL)", "Target\n(ng/mL)", "PRD Expected\nvalues (ng/mL)"],
            ["SCI 1", "21.3", "14.9 - 27.7", "21.3", "14.9 - 27.7", "21.6", "15.1 - 28.1"],
            ["SCI 2", "33.3", "23.3 - 43.3", "34.1", "23.9 - 44.3", "34.2", "23.9 - 44.5"],
            ["SCI 3", "42.2", "25.7 - 58.6", "43.4", "26.5 - 60.3", "43.3", "26.4 - 60.2"],
            ["SCI 4", "64.9", "39.6 - 90.3", "66.1", "40.3 - 91.9", "65.3", "39.8 - 90.8"],
            ["SCI 5", "82.5", "50.3 - 114.6", "85.1", "51.9 - 118.3", "85.1", "51.9 - 118.3"],
        ],
        [1.5, 2.1, 3.2, 2.1, 3.2, 2.1, 3.2],
        highlight_cells={(3, 2), (4, 4), (6, 1), (6, 4)},
        header_rows=1,
    )

    para(doc, "CONFIDENTIAL", bold=True, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(0)
    foot.paragraph_format.space_after = Pt(0)
    add_run(foot, "bioMérieux - Confidential Information", size=7.5)
    add_run(foot, "\t\t\t\t\tPage 12/46", size=7.5)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
